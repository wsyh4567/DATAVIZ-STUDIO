"""
将Markdown文档转换为格式美观的DOCX商业方案文档
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import re

def create_styled_document():
    """创建带有商业格式样式的文档"""
    doc = Document()
    
    # 设置文档默认字体为中文友好字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(10.5)
    
    # 设置页边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    return doc

def add_title(doc, text):
    """添加标题"""
    title = doc.add_heading(text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.name = '微软雅黑'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return title

def add_heading_1(doc, text):
    """添加一级标题"""
    heading = doc.add_heading(text, level=1)
    heading_run = heading.runs[0]
    heading_run.font.name = '微软雅黑'
    heading_run.font.size = Pt(16)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(0, 51, 102)
    heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return heading

def add_heading_2(doc, text):
    """添加二级标题"""
    heading = doc.add_heading(text, level=2)
    heading_run = heading.runs[0]
    heading_run.font.name = '微软雅黑'
    heading_run.font.size = Pt(14)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(31, 78, 121)
    heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return heading

def add_heading_3(doc, text):
    """添加三级标题"""
    heading = doc.add_heading(text, level=3)
    heading_run = heading.runs[0]
    heading_run.font.name = '微软雅黑'
    heading_run.font.size = Pt(12)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(68, 114, 196)
    heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """添加段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(10.5)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    return para

def add_bullet_point(doc, text, level=0):
    """添加项目符号列表"""
    para = doc.add_paragraph(text, style='List Bullet')
    para.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = para.runs[0]
    run.font.name = '微软雅黑'
    run.font.size = Pt(10.5)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return para

def add_table_from_markdown(doc, lines):
    """从Markdown表格创建Word表格"""
    # 解析表格行
    rows = []
    for line in lines:
        if '|' in line and not line.strip().startswith('|---'):
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                rows.append(cells)
    
    if not rows:
        return
    
    # 创建表格
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    
    # 填充表格
    for i, row_data in enumerate(rows):
        for j, cell_data in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_data
            # 设置字体
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(9)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            # 表头加粗
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                cell._element.get_or_add_tcPr().append(
                    parse_xml(r'<w:shd {} w:fill="4472C4"/>'.format(nsdecls('w')))
                )
    
    doc.add_paragraph()  # 表格后添加空行
    return table

def parse_xml(xml_string):
    """解析XML字符串"""
    from docx.oxml import parse_xml as docx_parse_xml
    return docx_parse_xml(xml_string)

def nsdecls(*prefixes):
    """命名空间声明"""
    from docx.oxml.ns import nsdecls as docx_nsdecls
    return docx_nsdecls(*prefixes)

def convert_markdown_to_docx(md_file, docx_file):
    """转换Markdown到DOCX"""
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    doc = create_styled_document()
    
    i = 0
    in_table = False
    table_lines = []
    in_code_block = False
    code_lines = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # 处理代码块
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                # 结束代码块
                in_code_block = False
                if code_lines:
                    para = doc.add_paragraph()
                    run = para.add_run('\n'.join(code_lines))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    para.paragraph_format.left_indent = Inches(0.5)
                    para.paragraph_format.space_before = Pt(6)
                    para.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # 处理表格
        if '|' in line and line.strip():
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # 表格结束
            add_table_from_markdown(doc, table_lines)
            in_table = False
            table_lines = []
        
        # 处理标题
        if line.startswith('# ') and not line.startswith('## '):
            add_title(doc, line[2:])
        elif line.startswith('## ') and not line.startswith('### '):
            add_heading_1(doc, line[3:])
        elif line.startswith('### ') and not line.startswith('#### '):
            add_heading_2(doc, line[4:])
        elif line.startswith('#### '):
            add_heading_3(doc, line[5:])
        
        # 处理分隔线
        elif line.startswith('---'):
            doc.add_page_break()
        
        # 处理列表
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            # 检测是否为粗体
            if text.startswith('**') and '**' in text[2:]:
                add_bullet_point(doc, text)
            else:
                add_bullet_point(doc, text)
        
        # 处理普通段落
        elif line.strip() and not line.startswith('|'):
            # 检测粗体
            if '**' in line:
                para = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = para.add_run(part[2:-2])
                        run.font.bold = True
                    else:
                        run = para.add_run(part)
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(10.5)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            else:
                add_paragraph(doc, line)
        
        i += 1
    
    # 保存文档
    doc.save(docx_file)
    print(f"文档已成功转换并保存到: {docx_file}")

if __name__ == "__main__":
    md_file = "服装销售预测方案_修订版.md"
    docx_file = "服装销售预测方案.docx"
    
    try:
        convert_markdown_to_docx(md_file, docx_file)
        print("转换完成！")
    except Exception as e:
        print(f"转换过程中出现错误: {e}")
        import traceback
        traceback.print_exc()
